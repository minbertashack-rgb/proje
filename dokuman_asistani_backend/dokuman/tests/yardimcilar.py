from __future__ import annotations

from dataclasses import dataclass
import importlib
from pathlib import Path
from typing import Iterable

from django.core.files import File
from docx import Document
from docx.shared import Pt
from dokuman.services.heading_parser import parse_document_structure

from .assets.mini_varlik_havuzu import VARLIK_HAVUZU


@dataclass(frozen=True)
class BelgeOgesi:
    tur: str
    metin: str
    seviye: int | None = None
    yazi_boyutu: float | None = None
    kalin: bool = False
    stil_baslik: bool = False
    yeni_sayfa: bool = False


def baslik(
    metin: str,
    *,
    seviye: int = 1,
    yazi_boyutu: float | None = None,
    kalin: bool = True,
    stil_baslik: bool = False,
    yeni_sayfa: bool = False,
) -> BelgeOgesi:
    return BelgeOgesi(
        tur="baslik",
        metin=metin,
        seviye=seviye,
        yazi_boyutu=yazi_boyutu,
        kalin=kalin,
        stil_baslik=stil_baslik,
        yeni_sayfa=yeni_sayfa,
    )


def paragraf(
    metin: str,
    *,
    yazi_boyutu: float | None = None,
    kalin: bool = False,
    yeni_sayfa: bool = False,
) -> BelgeOgesi:
    return BelgeOgesi(
        tur="paragraf",
        metin=metin,
        yazi_boyutu=yazi_boyutu,
        kalin=kalin,
        yeni_sayfa=yeni_sayfa,
    )


def _oge_kaydi_coz(oge_kaydi: dict) -> BelgeOgesi:
    if oge_kaydi.get("tur") == "baslik":
        return baslik(
            oge_kaydi["metin"],
            seviye=oge_kaydi.get("seviye", 1),
            yazi_boyutu=oge_kaydi.get("yazi_boyutu"),
            kalin=oge_kaydi.get("kalin", True),
            stil_baslik=oge_kaydi.get("stil_baslik", False),
            yeni_sayfa=oge_kaydi.get("yeni_sayfa", False),
        )

    return paragraf(
        oge_kaydi["metin"],
        yazi_boyutu=oge_kaydi.get("yazi_boyutu"),
        kalin=oge_kaydi.get("kalin", False),
        yeni_sayfa=oge_kaydi.get("yeni_sayfa", False),
    )


def fixture_docx_olustur(tmp_path: Path, dosya_adi: str, ogeler: Iterable[BelgeOgesi]) -> Path:
    belge = Document()

    for idx, oge in enumerate(ogeler):
        if oge.yeni_sayfa and idx > 0:
            belge.add_page_break()

        p = belge.add_paragraph()
        if oge.tur == "baslik" and oge.stil_baslik and oge.seviye:
            try:
                p.style = f"Heading {max(1, min(int(oge.seviye), 9))}"
            except Exception:
                pass

        run = p.add_run(oge.metin)
        run.bold = oge.kalin

        if oge.yazi_boyutu is not None:
            run.font.size = Pt(float(oge.yazi_boyutu))
        elif oge.tur == "baslik":
            run.font.size = Pt(16 if (oge.seviye or 1) <= 1 else 14)
        else:
            run.font.size = Pt(12)

    yol = tmp_path / dosya_adi
    belge.save(str(yol))
    return yol


def belgeyi_uret_docx(tmp_path: Path, dosya_adi: str, ogeler: Iterable[BelgeOgesi]) -> Path:
    return fixture_docx_olustur(tmp_path, dosya_adi, ogeler)


def fixture_pdf_olustur(tmp_path: Path, dosya_adi: str, ogeler: Iterable[BelgeOgesi]) -> Path:
    fitz = importlib.import_module("fitz")
    belge = fitz.open()
    sayfa = belge.new_page()
    y = 72

    for oge in ogeler:
        if oge.yeni_sayfa or y > 760:
            sayfa = belge.new_page()
            y = 72

        yazi_boyutu = oge.yazi_boyutu
        if yazi_boyutu is None:
            if oge.tur == "baslik":
                yazi_boyutu = 16 if (oge.seviye or 1) <= 1 else 14
            else:
                yazi_boyutu = 12

        sayfa.insert_text(
            (72, y),
            oge.metin,
            fontsize=float(yazi_boyutu),
            fontname="helv",
        )
        y += 28 if oge.tur == "baslik" else 22

    yol = tmp_path / dosya_adi
    belge.save(str(yol))
    belge.close()
    return yol


def belgeyi_uret_pdf(tmp_path: Path, dosya_adi: str, ogeler: Iterable[BelgeOgesi]) -> Path:
    return fixture_pdf_olustur(tmp_path, dosya_adi, ogeler)


def belge_fixture_yolu(tmp_path: Path, varlik_adi: str) -> Path:
    varlik = VARLIK_HAVUZU[varlik_adi]
    ogeler = [_oge_kaydi_coz(oge) for oge in varlik["ogeler"]]

    if varlik["tur"] == "pdf":
        return fixture_pdf_olustur(tmp_path, varlik["dosya_adi"], ogeler)
    if varlik["tur"] == "docx":
        return fixture_docx_olustur(tmp_path, varlik["dosya_adi"], ogeler)

    raise ValueError(f"Bilinmeyen varlik turu: {varlik['tur']}")


def pdf_destegi_var_mi() -> bool:
    try:
        importlib.import_module("fitz")
        return True
    except Exception:
        return False


def belge_yapisini_coz(dosya_yolu: Path) -> dict:
    return parse_document_structure(dosya_yolu)


def bolum_listesini_al(parsed: dict | list[dict]) -> list[dict]:
    if isinstance(parsed, dict):
        return list(parsed.get("sections") or [])
    return list(parsed)


def section_basliklarini_topla(parsed: dict | list[dict]) -> list[str]:
    return [sec.get("title") or "" for sec in bolum_listesini_al(parsed)]


def basliklari_getir(parsed: dict | list[dict]) -> list[str]:
    return section_basliklarini_topla(parsed)


def section_pathlerini_topla(parsed: dict | list[dict]) -> list[str]:
    return [sec.get("path") or "" for sec in bolum_listesini_al(parsed)]


def parca_adreslerini_topla(parcalar) -> list[str]:
    return [getattr(parca, "adres", "") for parca in parcalar]


def parcalari_getir(dokuman) -> list:
    return list(dokuman.parcalar.order_by("sira"))


def beklenen_basliklari_kontrol_et(parsed: dict | list[dict], beklenen_basliklar: list[str]) -> bool:
    return section_basliklarini_topla(parsed) == beklenen_basliklar


def adresler_benzersiz_mi(adresler: Iterable[str]) -> bool:
    temiz_adresler = [adres for adres in adresler if adres]
    return len(temiz_adresler) == len(set(temiz_adresler))


def parca_sayisi_makul_mu(parca_sayisi: int, *, min_sayi: int = 1, max_sayi: int | None = None) -> bool:
    if parca_sayisi < min_sayi:
        return False
    if max_sayi is not None and parca_sayisi > max_sayi:
        return False
    return True


def dokuman_kaydi_olustur(
    *,
    kullanici,
    dosya_yolu: Path,
    baslik: str,
    mime: str = "",
) -> "Dokuman":
    from dokuman.models import Dokuman

    with dosya_yolu.open("rb") as fh:
        django_file = File(fh, name=dosya_yolu.name)
        dokuman = Dokuman(
            owner=kullanici,
            baslik=baslik,
            mime=mime,
            durum="yuklendi",
            hata="",
        )
        dokuman.dosya.save(dosya_yolu.name, django_file, save=False)
        dokuman.save()
    return dokuman
