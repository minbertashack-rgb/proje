from dokuman.services import heading_parser
from django.test import override_settings


TextBlock = heading_parser.TextBlock


def make_block(text: str, **kwargs) -> TextBlock:
    defaults = {
        "page": 1,
        "font_size": 12.0,
        "bold": False,
    }
    defaults.update(kwargs)
    return TextBlock(text=text, **defaults)


def test_strict_special_headings_are_detected():
    samples = [
        "GIRIS",
        "GIRIS:",
        "GIRIS - Kapsam",
        "EK",
        "EK B",
        "EK-10",
        "ÖZET",
        "SONUÇ",
        "GİRİŞ",
        "SONUC - Degerlendirme",
    ]

    for text in samples:
        assert heading_parser._is_strict_special_heading(text) is True
        assert heading_parser.numbered_depth(text) == 1
        assert heading_parser.resolved_heading_level(make_block(text)) == 1
        assert heading_parser.is_heading(make_block(text), 12.0) is True


def test_sentence_like_text_is_not_misclassified_as_heading():
    samples = [
        "Ek notlar paragrafi.",
        "Giris paragrafi.",
        "Bu bir giris cumlesidir.",
    ]

    for text in samples:
        block = make_block(text)
        assert heading_parser._is_strict_special_heading(text) is False
        assert heading_parser.numbered_depth(text) is None
        assert heading_parser.resolved_heading_level(block) is None
        assert heading_parser.is_heading(block, 12.0) is False
        assert heading_parser.heading_score(block, 12.0) < 0.65


def test_multiline_heading_blocks_are_merged_when_visual_style_matches():
    blocks = [
        make_block("1. Giris", font_size=14.0, bold=True),
        make_block("ve Amac", font_size=14.0, bold=True),
        make_block("Bu ilk paragraftir.", font_size=12.0, bold=False),
    ]

    merged = heading_parser.merge_multiline_heading_blocks(blocks)

    assert len(merged) == 2
    assert merged[0].text == "1. Giris ve Amac"
    assert merged[0].bold is True
    assert merged[0].font_size == 14.0


def test_build_section_tree_puts_intro_text_before_first_heading():
    blocks = [
        make_block("Belge giris satiri.", font_size=12.0),
        make_block("1. Giris", font_size=14.0, bold=True),
        make_block("Birinci bolum icerigi.", font_size=12.0),
        make_block("2. Yontem", font_size=14.0, bold=True),
        make_block("Ikinci bolum icerigi.", font_size=12.0),
    ]

    root = heading_parser.build_section_tree(blocks)
    sections = heading_parser.flatten_sections(root)

    assert sections[0]["title"] == "Belge Başlangıcı"
    assert sections[0]["path"] == "0"
    assert sections[0]["content"] == "Belge giris satiri."
    assert sections[1]["title"] == "1. Giris"
    assert sections[1]["path"] == "1"
    assert sections[1]["content"] == "Birinci bolum icerigi."
    assert sections[2]["title"] == "2. Yontem"
    assert sections[2]["content"] == "Ikinci bolum icerigi."


def test_stress_intro_then_numbered_heading_paths_remain_distinct():
    blocks = [
        make_block("Belge giris metni.", font_size=12.0),
        make_block("1. Giris", font_size=14.0, bold=True),
        make_block("Birinci bolum icerigi.", font_size=12.0),
        make_block("1.1 Amac", font_size=13.0, bold=True),
        make_block("Amac metni.", font_size=12.0),
    ]

    root = heading_parser.build_section_tree(blocks)
    sections = heading_parser.flatten_sections(root)

    assert sections[0]["title"] == "Belge Başlangıcı"
    assert sections[0]["path"] == "0"
    assert sections[1]["title"] == "1. Giris"
    assert sections[1]["path"] == "1"
    assert sections[2]["title"] == "1.1 Amac"
    assert sections[2]["path"] == "1.1"


def test_all_caps_sentence_paragraph_is_not_misclassified_as_heading():
    text = "BU BOLUMDE SISTEMIN NASIL CALISTIGI DETAYLI OLARAK ANLATILMAKTADIR."
    block = make_block(text, font_size=12.0, bold=False)

    assert heading_parser.is_heading(block, 12.0) is False
    assert heading_parser.heading_score(block, 12.0) < 0.65


def test_all_caps_explanatory_line_without_period_is_not_heading_even_if_emphasized():
    text = "AMA VE KAPSAM BU BELGENIN TEMEL HEDEFLERINI ACIKLAR"
    block = make_block(text, font_size=13.0, bold=True)

    assert heading_parser.is_heading(block, 12.0) is False
    assert heading_parser.heading_score(block, 12.0) < 0.65


def test_broken_ocr_letter_soup_is_not_misclassified_as_heading():
    text = "B U T U N L E S I K A K I S I"
    block = make_block(text, font_size=13.2, bold=True)

    assert heading_parser.looks_like_broken_ocr_line(text) is True
    assert heading_parser.is_heading(block, 12.0) is False
    assert heading_parser.heading_score(block, 12.0) < 0.65


def test_regular_short_heading_is_not_blocked_by_ocr_noise_guard():
    text = "JWT Akisi"
    block = make_block(text, font_size=13.2, bold=True)

    assert heading_parser.looks_like_broken_ocr_line(text) is False
    assert heading_parser.is_heading(block, 12.0) is True


def test_real_short_special_heading_is_preserved():
    blocks = [
        make_block("Sonuc", font_size=14.0, bold=True),
        make_block("Kisa ama gercek sonuc paragrafi.", font_size=12.0),
    ]

    root = heading_parser.build_section_tree(blocks)
    sections = heading_parser.flatten_sections(root)

    assert len(sections) == 1
    assert sections[0]["title"] == "Sonuc"
    assert sections[0]["level"] == 1
    assert sections[0]["path"] == "1"
    assert sections[0]["content"] == "Kisa ama gercek sonuc paragrafi."


def test_single_word_real_heading_is_preserved():
    blocks = [
        make_block("Amaç", font_size=13.0, bold=True),
        make_block("Bu bolum kisa ama gercek bir amac aciklamasi verir.", font_size=12.0),
    ]

    root = heading_parser.build_section_tree(blocks)
    sections = heading_parser.flatten_sections(root)

    assert len(sections) == 1
    assert sections[0]["title"] == "Amaç"
    assert sections[0]["path"] == "1"
    assert sections[0]["content"] == "Bu bolum kisa ama gercek bir amac aciklamasi verir."


def test_short_clean_unnumbered_heading_with_slightly_larger_font_is_kept():
    blocks = [
        make_block("Yontem", font_size=13.0, bold=False),
        make_block("Bu satir ilgili yontem adimlarini kisa ve net sekilde aciklar.", font_size=12.0),
    ]

    root = heading_parser.build_section_tree(blocks)
    sections = heading_parser.flatten_sections(root)

    assert len(sections) == 1
    assert sections[0]["title"] == "Yontem"
    assert sections[0]["path"] == "1"
    assert sections[0]["content"] == "Bu satir ilgili yontem adimlarini kisa ve net sekilde aciklar."


def test_numbered_hierarchy_keeps_levels_and_paths_consistent():
    blocks = [
        make_block("1. Giris", font_size=15.0, bold=True),
        make_block("Giris icerigi.", font_size=12.0),
        make_block("1.1 Amac", font_size=14.0, bold=True),
        make_block("Amac icerigi.", font_size=12.0),
        make_block("1.2 Kapsam", font_size=14.0, bold=True),
        make_block("Kapsam icerigi.", font_size=12.0),
        make_block("2. Yontem", font_size=15.0, bold=True),
        make_block("Yontem icerigi.", font_size=12.0),
    ]

    root = heading_parser.build_section_tree(blocks)
    sections = heading_parser.flatten_sections(root)

    assert [s["title"] for s in sections] == [
        "1. Giris",
        "1.1 Amac",
        "1.2 Kapsam",
        "2. Yontem",
    ]
    assert [s["level"] for s in sections] == [1, 2, 2, 1]
    assert [s["path"] for s in sections] == ["1", "1.1", "1.2", "2"]


def test_appendix_like_heading_is_detected():
    block = make_block("EK A: Veri Seti", font_size=12.0, bold=False)

    assert heading_parser.is_heading(block, 12.0) is True
    assert heading_parser.resolved_heading_level(block) == 1


def test_long_sentence_ending_with_period_is_not_heading():
    text = (
        "Bu paragraf sistemin veri toplama, isleme ve raporlama adimlarini "
        "detayli bir sekilde aciklamaktadir."
    )
    block = make_block(text, font_size=12.0, bold=False)

    assert heading_parser.is_heading(block, 12.0) is False
    assert heading_parser.heading_score(block, 12.0) < 0.65


def test_long_heading_like_line_without_numbering_is_not_heading():
    text = (
        "Veri toplama ve model egitimi akisinda kullanilan temel kontrol "
        "adimlari ekip ici aciklama satiri olarak yazilmistir"
    )
    block = make_block(text, font_size=14.0, bold=True)

    assert heading_parser.is_heading(block, 12.0) is False
    assert heading_parser.heading_score(block, 12.0) < 0.65


def test_bold_sentence_ending_with_period_is_not_heading_even_if_emphasized():
    block = make_block(
        "Bu bolum sistemin risklerini kisa ama normal bir cumleyle aciklar.",
        font_size=13.4,
        bold=True,
    )

    assert heading_parser.is_heading(block, 12.0) is False
    assert heading_parser.heading_score(block, 12.0) < 0.65


def test_short_explanatory_line_without_period_is_not_heading():
    block = make_block(
        "Bu bolum temel riskleri ve kontrol mantigini aciklar",
        font_size=13.2,
        bold=True,
    )

    assert heading_parser.is_heading(block, 12.0) is False
    assert heading_parser.heading_score(block, 12.0) < 0.65


def test_real_short_heading_like_degerlendirme_is_preserved():
    block = make_block("Degerlendirme", font_size=13.5, bold=True)

    assert heading_parser.is_heading(block, 12.0) is True
    assert heading_parser.heading_score(block, 12.0) >= 0.65


def test_heading_score_details_exposes_safe_reason_for_real_short_heading():
    detay = heading_parser.heading_score_details(
        make_block("Amaç", font_size=13.2, bold=True),
        12.0,
    )

    assert detay["heading_score"] >= 0.65
    assert detay["heading_score"] <= 1.0
    assert detay["heading_decision_reason"] in {
        "bold_heading",
        "clean_short_heading",
        "short_valid_heading",
        "larger_font",
    }
    assert "heading_positive_signals" in detay
    assert "heading_penalty_signals" in detay


def test_heading_score_details_penalizes_short_sentence_false_positive():
    detay = heading_parser.heading_score_details(
        make_block("Bu bir uyaridir.", font_size=12.8, bold=False),
        12.0,
    )

    assert detay["heading_score"] < 0.65
    assert detay["heading_decision_reason"] in {
        "sentence_ending_penalty",
        "plain_sentence_penalty",
        "short_sentence_false_positive_penalty",
    }


def test_heading_regression_matrix_covers_short_valid_and_false_positive_cases():
    vakalar = [
        ("Amaç", True, 13.2, True),
        ("Tanım", True, 13.0, True),
        ("1. Giris", True, 14.0, True),
        ("1.1 Kapsam", True, 13.4, True),
        ("EK A: Veri Seti", True, 12.0, False),
        ("Bu bir giris cumlesidir.", False, 12.0, False),
        ("UYARI: BU BELGEDE EK ACIKLAMA VARDIR", False, 13.0, True),
        ("Kisa not:", False, 12.4, False),
        ("AMA VE KAPSAM BU BELGENIN TEMEL HEDEFLERINI ACIKLAR", False, 13.0, True),
        ("Sonuç", True, 13.4, True),
    ]

    for text, beklenen, font_size, bold in vakalar:
        block = make_block(text, font_size=font_size, bold=bold)
        assert heading_parser.is_heading(block, 12.0) is beklenen, text


def test_parse_document_structure_adds_safe_heading_debug_meta_when_enabled(tmp_path):
    dosya = tmp_path / "debug-heading.docx"
    from .yardimcilar import belgeyi_uret_docx, baslik, paragraf

    belgeyi_uret_docx(
        tmp_path,
        "debug-heading.docx",
        [
            baslik("Amaç", seviye=1, yazi_boyutu=14, kalin=True),
            paragraf("Bu bolum belge amacini aciklar.", yazi_boyutu=12),
        ],
    )


def test_docx_table_rows_are_preserved_in_section_content(tmp_path):
    from docx import Document

    dosya = tmp_path / "table-section.docx"
    doc = Document()
    doc.add_heading("Veri Ozeti", level=1)
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Alan"
    table.cell(0, 1).text = "Deger"
    table.cell(1, 0).text = "JWT"
    table.cell(1, 1).text = "Token"
    doc.save(str(dosya))

    parsed = heading_parser.parse_document_structure(dosya)
    sections = parsed["sections"]

    assert sections[0]["title"] == "Veri Ozeti"
    assert "Tablo 1 Satir 1" in sections[0]["content"]
    assert "JWT | Token" in sections[0]["content"]

    with override_settings(DOCVERSE_DEBUG_SUMMARY_ENABLED=True):
        parsed = heading_parser.parse_document_structure(dosya)

    assert "debug_ozeti" in parsed
    assert parsed["sections"][0]["heading_score"] >= 0.65
    assert parsed["sections"][0]["heading_decision_reason"] != ""
    assert "content" in parsed["sections"][0]
    assert "heading_reason_ozeti" in parsed["debug_ozeti"]
