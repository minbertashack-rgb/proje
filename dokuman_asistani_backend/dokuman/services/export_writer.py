from __future__ import annotations

from io import BytesIO
from pathlib import Path
import textwrap
from typing import Any
import xml.etree.ElementTree as ET
from xml.sax.saxutils import escape as xml_escape
import zipfile

from django.conf import settings


SOURCE_MANIFEST_VERSION = "2.0"


def _clean_text(value: Any) -> str:
    return " ".join(str(value or "").split()).strip()


def _slug(value: str, *, fallback: str) -> str:
    clean = "".join(ch.lower() if ch.isalnum() else "_" for ch in _clean_text(value))
    clean = "_".join(part for part in clean.split("_") if part)
    return (clean[:48] or fallback).strip("_") or fallback


def _manifest_source_ids(manifest: dict | None) -> list[int]:
    seen = set()
    out = []
    for value in (manifest or {}).get("kaynak_parca_idleri") or []:
        try:
            clean = int(value)
        except Exception:
            continue
        if clean in seen:
            continue
        seen.add(clean)
        out.append(clean)
    return out


def build_output_meta(
    *,
    manifest: dict | None,
    output_format: str,
    output_created: bool,
    section_count: int = 0,
    slide_count: int = 0,
) -> dict:
    source_ids = _manifest_source_ids(manifest)
    return {
        "output_format": _clean_text(output_format).lower(),
        "source_manifest_version": SOURCE_MANIFEST_VERSION,
        "source_parca_idleri": source_ids,
        "section_count": int(section_count or 0),
        "slide_count": int(slide_count or 0),
        "output_created": bool(output_created),
    }


def build_attachment_filename(*, title: str, prefix: str, ext: str) -> str:
    stem = _slug(title, fallback=prefix)
    return f"{prefix}_{stem}.{ext.lstrip('.')}"


def build_markdown_document(*, title: str, sections: list[dict], output_meta: dict | None = None) -> str:
    lines = [f"# {title}", ""]
    for section in sections or []:
        heading = _clean_text(section.get("heading") or section.get("title"))
        if heading:
            lines.append(f"## {heading}")
        body = section.get("body")
        bullets = list(section.get("bullets") or [])
        if body:
            lines.append(str(body))
        if bullets:
            for item in bullets:
                clean = _clean_text(item)
                if clean:
                    lines.append(f"- {clean}")
        lines.append("")

    if output_meta:
        lines.extend(
            [
                "## Output Meta",
                f"- output_format: {output_meta.get('output_format')}",
                f"- source_manifest_version: {output_meta.get('source_manifest_version')}",
                f"- source_parca_idleri: {', '.join(str(item) for item in output_meta.get('source_parca_idleri') or [])}",
                f"- section_count: {output_meta.get('section_count')}",
                f"- slide_count: {output_meta.get('slide_count')}",
                f"- output_created: {output_meta.get('output_created')}",
                "",
            ]
        )
    return "\n".join(lines).strip() + "\n"


def markdown_to_txt(markdown_text: str) -> str:
    return (
        str(markdown_text or "")
        .replace("# ", "")
        .replace("## ", "")
        .replace("### ", "")
        .replace("**", "")
    )


def write_docx_bytes(*, title: str, sections: list[dict]) -> tuple[bytes | None, str]:
    try:
        from docx import Document
    except Exception:
        return None, "docx_runtime_unavailable"

    document = Document()
    document.add_heading(title, level=1)
    for section in sections or []:
        heading = _clean_text(section.get("heading") or section.get("title"))
        if heading:
            document.add_heading(heading, level=2)
        body = _clean_text(section.get("body"))
        if body:
            document.add_paragraph(body)
        for item in section.get("bullets") or []:
            clean = _clean_text(item)
            if clean:
                document.add_paragraph(clean, style="List Bullet")

    bio = BytesIO()
    document.save(bio)
    return bio.getvalue(), ""


def _pdf_escape(value: str) -> str:
    clean = _clean_text(value)
    clean = clean.encode("latin-1", errors="replace").decode("latin-1")
    return clean.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _pdf_lines(*, title: str, sections: list[dict]) -> list[str]:
    lines = [title, ""]
    for section in sections or []:
        heading = _clean_text(section.get("heading") or section.get("title"))
        if heading:
            lines.append(heading.upper())
        body = _clean_text(section.get("body"))
        if body:
            lines.extend(textwrap.wrap(body, width=88) or [body])
        for item in section.get("bullets") or []:
            clean = _clean_text(item)
            if clean:
                wrapped = textwrap.wrap(f"- {clean}", width=88) or [f"- {clean}"]
                lines.extend(wrapped)
        lines.append("")
    return lines or ["Bos export."]


def _build_pdf_page_stream(lines: list[str]) -> bytes:
    content = ["BT", "/F1 11 Tf", "50 800 Td", "14 TL"]
    first_line = True
    for line in lines:
        text = _pdf_escape(line)
        if first_line:
            content.append(f"({text}) Tj")
            first_line = False
        else:
            content.append("T*")
            content.append(f"({text}) Tj")
    content.append("ET")
    return "\n".join(content).encode("latin-1", errors="replace")


def write_pdf_bytes(*, title: str, sections: list[dict]) -> tuple[bytes | None, str]:
    lines = _pdf_lines(title=title, sections=sections)
    page_size = 46
    pages = [lines[idx : idx + page_size] for idx in range(0, len(lines), page_size)] or [["Bos export."]]
    font_id = 1
    page_count = len(pages)
    content_start = 2
    page_start = content_start + page_count
    pages_id = page_start + page_count
    catalog_id = pages_id + 1

    objects: dict[int, bytes] = {
        font_id: b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    }
    for idx, page_lines in enumerate(pages):
        stream = _build_pdf_page_stream(page_lines)
        objects[content_start + idx] = (
            f"<< /Length {len(stream)} >>\nstream\n".encode("latin-1")
            + stream
            + b"\nendstream"
        )
    for idx in range(page_count):
        objects[page_start + idx] = (
            f"<< /Type /Page /Parent {pages_id} 0 R /MediaBox [0 0 595 842] "
            f"/Resources << /Font << /F1 {font_id} 0 R >> >> "
            f"/Contents {content_start + idx} 0 R >>"
        ).encode("latin-1")
    kids = " ".join(f"{page_start + idx} 0 R" for idx in range(page_count))
    objects[pages_id] = f"<< /Type /Pages /Kids [{kids}] /Count {page_count} >>".encode("latin-1")
    objects[catalog_id] = f"<< /Type /Catalog /Pages {pages_id} 0 R >>".encode("latin-1")

    output = BytesIO()
    output.write(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    xref_positions = [0]
    for object_id in range(1, catalog_id + 1):
        xref_positions.append(output.tell())
        output.write(f"{object_id} 0 obj\n".encode("latin-1"))
        output.write(objects[object_id])
        output.write(b"\nendobj\n")
    xref_start = output.tell()
    output.write(f"xref\n0 {catalog_id + 1}\n".encode("latin-1"))
    output.write(b"0000000000 65535 f \n")
    for pos in xref_positions[1:]:
        output.write(f"{pos:010d} 00000 n \n".encode("latin-1"))
    output.write(
        (
            f"trailer\n<< /Size {catalog_id + 1} /Root {catalog_id} 0 R >>\n"
            f"startxref\n{xref_start}\n%%EOF"
        ).encode("latin-1")
    )
    return output.getvalue(), ""


def _pptx_template_path() -> Path | None:
    base_dir = Path(getattr(settings, "BASE_DIR", Path.cwd()))
    candidates = [
        base_dir / "test_ingest.pptx",
        base_dir / "media" / "dokuman_asistani" / "test_ingest.pptx",
    ]
    return next((path for path in candidates if path.exists()), None)


def _pptx_text_paragraphs(items: list[str]) -> str:
    paragraphs = []
    for item in items or ["Bos slayt."]:
        clean = xml_escape(_clean_text(item) or "Bos slayt.")
        paragraphs.append(
            "<a:p><a:pPr lvl=\"0\"/><a:r><a:rPr lang=\"en-US\" sz=\"1800\"/>"
            f"<a:t>{clean}</a:t></a:r></a:p>"
        )
    return "".join(paragraphs)


def _build_slide_xml(*, title: str, bullets: list[str], notes: str) -> str:
    title_text = xml_escape(_clean_text(title) or "Slayt")
    body_items = [item for item in bullets if _clean_text(item)]
    if notes and len(body_items) < 5:
        body_items.append(f"Konusma notu: {notes}")
    return (
        "<?xml version='1.0' encoding='UTF-8' standalone='yes'?>"
        "<p:sld xmlns:a=\"http://schemas.openxmlformats.org/drawingml/2006/main\" "
        "xmlns:p=\"http://schemas.openxmlformats.org/presentationml/2006/main\" "
        "xmlns:r=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships\">"
        "<p:cSld><p:spTree>"
        "<p:nvGrpSpPr><p:cNvPr id=\"1\" name=\"\"/><p:cNvGrpSpPr/><p:nvPr/></p:nvGrpSpPr><p:grpSpPr/>"
        "<p:sp><p:nvSpPr><p:cNvPr id=\"2\" name=\"Title 1\"/><p:cNvSpPr><a:spLocks noGrp=\"1\"/></p:cNvSpPr>"
        "<p:nvPr><p:ph type=\"title\"/></p:nvPr></p:nvSpPr><p:spPr/>"
        "<p:txBody><a:bodyPr/><a:lstStyle/>"
        f"<a:p><a:r><a:t>{title_text}</a:t></a:r></a:p></p:txBody></p:sp>"
        "<p:sp><p:nvSpPr><p:cNvPr id=\"3\" name=\"TextBox 2\"/><p:cNvSpPr txBox=\"1\"/><p:nvPr/></p:nvSpPr>"
        "<p:spPr><a:xfrm><a:off x=\"457200\" y=\"1600200\"/><a:ext cx=\"8229600\" cy=\"4114800\"/></a:xfrm>"
        "<a:prstGeom prst=\"rect\"><a:avLst/></a:prstGeom><a:noFill/></p:spPr>"
        "<p:txBody><a:bodyPr wrap=\"square\"><a:spAutoFit/></a:bodyPr><a:lstStyle/>"
        f"{_pptx_text_paragraphs(body_items)}</p:txBody></p:sp>"
        "</p:spTree></p:cSld><p:clrMapOvr><a:masterClrMapping/></p:clrMapOvr></p:sld>"
    )


def _update_presentation_xml(raw_xml: bytes, slide_count: int) -> bytes:
    ns = {
        "p": "http://schemas.openxmlformats.org/presentationml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    ET.register_namespace("a", "http://schemas.openxmlformats.org/drawingml/2006/main")
    ET.register_namespace("p", ns["p"])
    ET.register_namespace("r", ns["r"])
    root = ET.fromstring(raw_xml)
    slide_list = root.find("p:sldIdLst", ns)
    if slide_list is None:
        slide_list = ET.SubElement(root, f"{{{ns['p']}}}sldIdLst")
    slide_list.clear()
    for idx in range(slide_count):
        ET.SubElement(
            slide_list,
            f"{{{ns['p']}}}sldId",
            {
                "id": str(256 + idx),
                f"{{{ns['r']}}}id": f"rId{7 + idx}",
            },
        )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _update_presentation_rels(raw_xml: bytes, slide_count: int) -> bytes:
    ns = {"rel": "http://schemas.openxmlformats.org/package/2006/relationships"}
    ET.register_namespace("", ns["rel"])
    root = ET.fromstring(raw_xml)
    for rel in list(root.findall("rel:Relationship", ns)):
        target = rel.attrib.get("Target", "")
        if target.startswith("slides/slide"):
            root.remove(rel)
    for idx in range(slide_count):
        ET.SubElement(
            root,
            f"{{{ns['rel']}}}Relationship",
            {
                "Id": f"rId{7 + idx}",
                "Type": "http://schemas.openxmlformats.org/officeDocument/2006/relationships/slide",
                "Target": f"slides/slide{idx + 1}.xml",
            },
        )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def _update_content_types(raw_xml: bytes, slide_count: int) -> bytes:
    ns = {"ct": "http://schemas.openxmlformats.org/package/2006/content-types"}
    ET.register_namespace("", ns["ct"])
    root = ET.fromstring(raw_xml)
    for node in list(root.findall("ct:Override", ns)):
        if node.attrib.get("PartName", "").startswith("/ppt/slides/slide"):
            root.remove(node)
    for idx in range(slide_count):
        ET.SubElement(
            root,
            f"{{{ns['ct']}}}Override",
            {
                "PartName": f"/ppt/slides/slide{idx + 1}.xml",
                "ContentType": "application/vnd.openxmlformats-officedocument.presentationml.slide+xml",
            },
        )
    return ET.tostring(root, encoding="utf-8", xml_declaration=True)


def write_pptx_bytes(*, title: str, slides: list[dict]) -> tuple[bytes | None, str]:
    template_path = _pptx_template_path()
    if template_path is None:
        return None, "pptx_template_missing"

    normalized_slides = list(slides or [])[:8]
    if not normalized_slides:
        normalized_slides = [{"title": title, "bullets": ["Sunum icerigi hazirlanamadi."], "notes": ""}]

    with zipfile.ZipFile(template_path, "r") as src, BytesIO() as bio:
        with zipfile.ZipFile(bio, "w", compression=zipfile.ZIP_DEFLATED) as dst:
            for name in src.namelist():
                if name in {
                    "[Content_Types].xml",
                    "ppt/presentation.xml",
                    "ppt/_rels/presentation.xml.rels",
                }:
                    continue
                if name.startswith("ppt/slides/slide") or name.startswith("ppt/slides/_rels/slide"):
                    continue
                dst.writestr(name, src.read(name))

            slide_rels = src.read("ppt/slides/_rels/slide1.xml.rels")
            for idx, slide in enumerate(normalized_slides, start=1):
                bullets = list(slide.get("bullets") or [])[:4]
                notes = _clean_text(slide.get("notes"))
                dst.writestr(
                    f"ppt/slides/slide{idx}.xml",
                    _build_slide_xml(
                        title=_clean_text(slide.get("title")) or f"Slayt {idx}",
                        bullets=bullets,
                        notes=notes,
                    ).encode("utf-8"),
                )
                dst.writestr(f"ppt/slides/_rels/slide{idx}.xml.rels", slide_rels)

            dst.writestr(
                "ppt/presentation.xml",
                _update_presentation_xml(src.read("ppt/presentation.xml"), len(normalized_slides)),
            )
            dst.writestr(
                "ppt/_rels/presentation.xml.rels",
                _update_presentation_rels(src.read("ppt/_rels/presentation.xml.rels"), len(normalized_slides)),
            )
            dst.writestr(
                "[Content_Types].xml",
                _update_content_types(src.read("[Content_Types].xml"), len(normalized_slides)),
            )
        return bio.getvalue(), ""


def apply_output_meta_headers(response, output_meta: dict):
    response["X-DocVerse-Output-Format"] = str(output_meta.get("output_format") or "")
    response["X-DocVerse-Source-Manifest-Version"] = str(output_meta.get("source_manifest_version") or "")
    response["X-DocVerse-Source-Part-Ids"] = ",".join(
        str(int(item))
        for item in (output_meta.get("source_parca_idleri") or [])[:12]
        if str(item).strip()
    )
    response["X-DocVerse-Section-Count"] = str(int(output_meta.get("section_count") or 0))
    response["X-DocVerse-Slide-Count"] = str(int(output_meta.get("slide_count") or 0))
    response["X-DocVerse-Output-Created"] = "1" if output_meta.get("output_created") else "0"
    return response
