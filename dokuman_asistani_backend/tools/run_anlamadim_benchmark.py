from __future__ import annotations

import argparse
import json
import os
from pathlib import Path
import socket
import sys
import tempfile
import time
import urllib.request
from urllib.parse import urlparse
from uuid import uuid4


BASE_DIR = Path(__file__).resolve().parent.parent
if str(BASE_DIR) not in sys.path:
    sys.path.insert(0, str(BASE_DIR))

os.environ.setdefault("DJANGO_SETTINGS_MODULE", "dokuman_asistani.settings")

import django

django.setup()

from django.contrib.auth import get_user_model
from django.test.utils import override_settings
from docx import Document
from rest_framework.test import APIClient

from dokuman.ai2.llm import ai2_istemcisini_hazirla
from dokuman.tests.anlamadim_yardimcilari import (
    benchmark_ornekleri,
    benchmark_toplam_rapor,
    genel_kalite_ozeti,
    okunur_benchmark_ozeti,
)


def _arg_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser(description="ParcaAnlamadimV2 kalite benchmark araci")
    parser.add_argument("--mode", choices=("smoke", "detailed"), default="smoke")
    parser.add_argument("--max-tokens", type=int, default=0)
    parser.add_argument("--sample-limit", type=int, default=0)
    parser.add_argument("--ai2-timeout", type=int, default=20)
    parser.add_argument("--preflight-timeout", type=int, default=20)
    parser.add_argument("--json-out", default=str(BASE_DIR / "anlamadim_benchmark_sonuc.json"))
    parser.add_argument("--text-out", default=str(BASE_DIR / "anlamadim_benchmark_ozet.txt"))
    return parser


def _ornek_docx_uret(hedef_klasor: Path, slug: str, baslik: str, paragraflar: list[str]) -> Path:
    belge = Document()
    try:
        p = belge.add_paragraph()
        p.style = "Heading 1"
        p.add_run(baslik).bold = True
    except Exception:
        run = belge.add_paragraph().add_run(baslik)
        run.bold = True

    for paragraf in paragraflar:
        belge.add_paragraph(paragraf)

    yol = hedef_klasor / f"{slug}.docx"
    belge.save(str(yol))
    return yol


def _benchmark_kullanicisi_olustur():
    kullanici_modeli = get_user_model()
    return kullanici_modeli.objects.create_user(
        username=f"anlamadim_benchmark_{uuid4().hex[:8]}",
        password="12345678",
    )


def _parca_sec(parcalar: list[dict]) -> dict | None:
    if not parcalar:
        return None
    for parca in parcalar:
        metin = str(parca.get("metin") or "").strip()
        if metin and len(metin) >= 30:
            return parca
    return parcalar[0]


def _ai2_preflight(timeout_saniye: int) -> tuple[bool, str]:
    istemci = ai2_istemcisini_hazirla()
    aday_adres = str((istemci.get("aday_adresler") or [""])[0] or "")
    parsed = urlparse(aday_adres)
    host = parsed.hostname or "127.0.0.1"
    port = int(parsed.port or (443 if parsed.scheme == "https" else 80))
    kok_adres = f"{parsed.scheme or 'http'}://{host}:{port}"
    models_adresi = f"{kok_adres}/v1/models"

    try:
        with socket.create_connection((host, port), timeout=max(1, int(timeout_saniye or 5))):
            with urllib.request.urlopen(models_adresi, timeout=max(1, int(timeout_saniye or 5))) as yanit:
                status = int(getattr(yanit, "status", 200) or 200)
            return True, f"AI2 preflight OK: {models_adresi} ({status})"
    except OSError as exc:
        return False, f"AI2 erisilemiyor: {host}:{port} ({exc})"
    except Exception as exc:
        return False, f"AI2 preflight basarisiz: {models_adresi} ({exc})"


def _tek_ornek_kos(client: APIClient, tmp_dir: Path, ornek: dict, max_tokens: int) -> dict:
    dosya_yolu = _ornek_docx_uret(tmp_dir, ornek["slug"], ornek["baslik"], list(ornek["paragraflar"]))

    with dosya_yolu.open("rb") as fh:
        yukle = client.post(
            "/api/dokuman-asistani/dokumanlar/yukle/",
            {"file": fh, "baslik": ornek["baslik"]},
            format="multipart",
        )

    if yukle.status_code != 201:
        return {
            "ornek_adi": ornek["slug"],
            "hata": f"upload_failed:{yukle.status_code}",
            "detay": dict(getattr(yukle, "data", {}) or {}),
        }

    doc_id = int((yukle.data or {}).get("id"))
    parcalar_resp = client.get(f"/api/dokuman-asistani/dokumanlar/{doc_id}/parcalar/")
    if parcalar_resp.status_code != 200:
        return {
            "ornek_adi": ornek["slug"],
            "doc_id": doc_id,
            "hata": f"parcalar_failed:{parcalar_resp.status_code}",
            "detay": dict(getattr(parcalar_resp, "data", {}) or {}),
        }

    parca = _parca_sec(list((parcalar_resp.data or {}).get("parcalar") or []))
    if not parca:
        return {
            "ornek_adi": ornek["slug"],
            "doc_id": doc_id,
            "hata": "parca_bulunamadi",
        }

    baslangic = time.perf_counter()
    anlamadim_resp = client.post(
        f"/api/dokuman-asistani/parcalar/{parca['id']}/anlamadim-v2/",
        {
            "mesaj": ornek.get("mesaj") or "Bu parcayi daha basit anlat.",
            "max_tokens": max_tokens,
            "debug_ai2": True,
        },
        format="json",
    )
    sure = round(time.perf_counter() - baslangic, 2)

    if anlamadim_resp.status_code != 200:
        return {
            "ornek_adi": ornek["slug"],
            "doc_id": doc_id,
            "parca_id": parca["id"],
            "hata": f"anlamadim_failed:{anlamadim_resp.status_code}",
            "detay": dict(getattr(anlamadim_resp, "data", {}) or {}),
            "sure_saniye": sure,
        }

    yanit = dict(anlamadim_resp.data or {})
    kalite = genel_kalite_ozeti(yanit, str(parca.get("metin") or ""), ornek_adi=ornek["slug"])
    kalite.update(
        {
            "doc_id": doc_id,
            "parca_id": int(parca["id"]),
            "adres": parca.get("adres") or "",
            "sure_saniye": sure,
            "dokumanda_yok": bool(yanit.get("dokumanda_yok")),
            "debug_ai2": dict(yanit.get("debug_ai2") or {}),
            "yanit": yanit,
        }
    )
    return kalite


def main() -> int:
    args = _arg_parser().parse_args()
    max_tokens = int(args.max_tokens) if int(args.max_tokens or 0) > 0 else (96 if args.mode == "smoke" else 240)

    json_out = Path(args.json_out)
    text_out = Path(args.text_out)
    json_out.parent.mkdir(parents=True, exist_ok=True)
    text_out.parent.mkdir(parents=True, exist_ok=True)

    preflight_ok, preflight_mesaji = _ai2_preflight(args.preflight_timeout)

    with tempfile.TemporaryDirectory() as tmp:
        tmp_path = Path(tmp)
        media_root = tmp_path / "media"
        media_root.mkdir(parents=True, exist_ok=True)

        kullanici = _benchmark_kullanicisi_olustur()
        client = APIClient()
        client.force_authenticate(user=kullanici)

        try:
            ornekler = benchmark_ornekleri(args.mode)
            if int(args.sample_limit or 0) > 0:
                ornekler = ornekler[: int(args.sample_limit)]

            with override_settings(
                MEDIA_ROOT=media_root,
                AI2_ZAMAN_ASIMI=max(5, int(args.ai2_timeout or 20)),
                YEREL_MODEL_ETKIN=False,
            ):
                sonuclar = []
                for ornek in ornekler:
                    sonuc = _tek_ornek_kos(client, tmp_path, ornek, max_tokens=max_tokens)
                    sonuclar.append(sonuc)

                tamamlananlar = [item for item in sonuclar if "hata" not in item]
                rapor = benchmark_toplam_rapor(tamamlananlar)
                rapor.update(
                    {
                        "mod": args.mode,
                        "max_tokens": max_tokens,
                        "sample_limit": int(args.sample_limit or 0),
                        "ai2_timeout": max(5, int(args.ai2_timeout or 20)),
                        "toplam_ornek": len(sonuclar),
                        "tamamlanan_ornek": len(tamamlananlar),
                        "ai2_preflight_ok": preflight_ok,
                        "ai2_preflight_mesaji": preflight_mesaji,
                        "basarisiz_ornekler": [item for item in sonuclar if "hata" in item],
                    }
                )

                json_out.write_text(json.dumps(rapor, ensure_ascii=False, indent=2), encoding="utf-8")
                text_ozet = okunur_benchmark_ozeti(rapor)
                text_ozet += f"\nAI2 Preflight: {preflight_mesaji}\n"
                if rapor.get("basarisiz_ornekler"):
                    text_ozet += "\nBasarisiz Ornekler:\n"
                    for item in rapor["basarisiz_ornekler"]:
                        text_ozet += f"- {item.get('ornek_adi')}: {item.get('hata')}\n"
                text_out.write_text(text_ozet, encoding="utf-8")

                print(text_ozet)
                print(f"JSON rapor: {json_out}")
                print(f"Yazi ozeti: {text_out}")

                if not tamamlananlar:
                    return 2
                return 0
        finally:
            kullanici.delete()


if __name__ == "__main__":
    raise SystemExit(main())
